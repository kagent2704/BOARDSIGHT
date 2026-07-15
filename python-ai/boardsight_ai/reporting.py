from __future__ import annotations

import csv
import json
from pathlib import Path

from boardsight_ai.models import PipelineResult


def _coalesce_mapping_value(mapping: dict[str, object], *keys: str, default: object = "n/a") -> object:
    for key in keys:
        if key in mapping and mapping[key] not in (None, ""):
            return mapping[key]
    return default


def write_structured_reports(result: PipelineResult, output_dir: Path) -> dict[str, str]:
    output_dir.mkdir(parents=True, exist_ok=True)
    files: dict[str, str] = {}

    transcript_csv = output_dir / "transcript.csv"
    with transcript_csv.open("w", newline="", encoding="utf-8") as handle:
        writer = csv.writer(handle)
        writer.writerow(["start", "end", "speaker", "text", "confidence"])
        for segment in result.transcript.segments:
            writer.writerow([segment.start, segment.end, segment.speaker, segment.text, segment.confidence])
    files["excel_ready_csv"] = str(transcript_csv)

    report_body = build_markdown_report(result)

    report_md = output_dir / "structured_report.md"
    report_md.write_text(report_body, encoding="utf-8")
    files["markdown_report"] = str(report_md)

    try:
        write_docx_report(result, output_dir / "structured_report.docx")
        files["docx"] = str(output_dir / "structured_report.docx")
    except Exception as exc:
        files["docx_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    try:
        write_pdf_report(result, output_dir / "structured_report.pdf")
        files["pdf"] = str(output_dir / "structured_report.pdf")
    except Exception as exc:
        files["pdf_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    try:
        write_excel_report(result, output_dir / "structured_report.xlsx")
        files["xlsx"] = str(output_dir / "structured_report.xlsx")
    except Exception as exc:
        files["xlsx_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    try:
        write_summary_image(result, output_dir / "summary_card.png")
        files["image"] = str(output_dir / "summary_card.png")
    except Exception as exc:
        files["image_error"] = f"{type(exc).__name__}: {str(exc).splitlines()[0]}"

    return files


def build_markdown_report(result: PipelineResult) -> str:
    cognitive_rating = result.meeting_scores.cognitive_rating or {}
    focus = _coalesce_mapping_value(cognitive_rating, "meeting_focus", "focus")
    clarity = _coalesce_mapping_value(cognitive_rating, "meeting_clarity", "clarity")
    overload_risk = _coalesce_mapping_value(cognitive_rating, "overload_risk")
    lines = [
        "# BoardSight Structured Report",
        "",
        "## Executive Summary",
        f"- Impact score: {result.meeting_scores.impact_score}",
        f"- Productivity score: {result.meeting_scores.productivity_score}",
        f"- Execution readiness: {result.meeting_scores.execution_readiness}",
        f"- Overall attention: {result.attention_sentiment.overall_attention}",
        f"- Overall sentiment: {result.attention_sentiment.overall_sentiment}",
        f"- Attention model sources: {', '.join(result.attention_sentiment.model_sources) or 'fallback'}",
        f"- Conclusion: {result.meeting_scores.meeting_conclusion}",
        "",
        "## Decision Moments",
    ]
    if result.decision_moments:
        for item in result.decision_moments:
            lines.append(
                f"- {item.event_id} at {item.timestamp}: {item.speaker} | {item.label} | {item.text}"
            )
    else:
        lines.append("- No explicit decision moments detected.")

    lines.extend(
        [
            "",
            "## Prioritized Decisions",
        ]
    )
    for item in result.workflow_model.prioritized_decisions:
        decision_id = item.get("decision_id", "unknown-decision")
        execution_rank = item.get("execution_rank", "?")
        priority_score = item.get("priority_score", "?")
        speaker = item.get("speaker", "Unknown speaker")
        text = item.get("text", item.get("summary", "No decision text available."))
        lines.append(
            f"- Rank {execution_rank}: {decision_id} | score {priority_score} | {speaker} | {text}"
        )

    lines.extend(
        [
            "",
            "## Execution Plan",
        ]
    )
    for task in result.workflow_model.execution_plan:
        execution_order = task.get("execution_order", "?")
        title = task.get("title", "Untitled task")
        owner = task.get("owner", "Unassigned")
        priority_score = task.get("priority_score", "?")
        lines.append(
            f"- {execution_order}. {title} | owner: {owner} | priority: {priority_score}"
        )

    lines.extend(
        [
            "",
            "## Visual Evidence",
        ]
    )
    for artifact in result.visual_artifacts[:10]:
        lines.append(
            f"- {artifact.artifact_id}: {artifact.artifact_type} | {artifact.display_mode} | {artifact.content_summary}"
        )

    lines.extend(
        [
            "",
            "## Speaker Ratings",
        ]
    )
    for speaker, rating in result.meeting_scores.speaker_rating.items():
        lines.append(f"- {speaker}: {rating}")

    lines.extend(
        [
            "",
            "## Cognitive Rating",
            f"- Focus: {focus}",
            f"- Clarity: {clarity}",
            f"- Overload risk: {overload_risk}",
            "",
            "## Organizational Impact Analysis",
            "- Operational quality: decisions with owners improve coordination and execution discipline.",
            "- Market standing: clear decisions reduce ambiguity in external-facing strategic execution.",
            "- Economic sensitivity: execution readiness and decision clarity influence cost of delay and rework risk.",
        ]
    )
    return "\n".join(lines) + "\n"


def write_docx_report(result: PipelineResult, path: Path) -> None:
    import docx

    document = docx.Document()
    document.add_heading("BoardSight Structured Report", level=1)
    document.add_paragraph(f"Impact score: {result.meeting_scores.impact_score}")
    document.add_paragraph(f"Productivity score: {result.meeting_scores.productivity_score}")
    document.add_paragraph(f"Execution readiness: {result.meeting_scores.execution_readiness}")
    document.add_paragraph(f"Conclusion: {result.meeting_scores.meeting_conclusion}")
    document.add_heading("Decision Moments", level=2)
    if result.decision_moments:
        for item in result.decision_moments:
            document.add_paragraph(
                f"{item.event_id} | {item.timestamp} | {item.speaker} | {item.label} | {item.text}",
                style="List Bullet",
            )
    else:
        document.add_paragraph("No explicit decision moments detected.")
    document.add_heading("Prioritized Decisions", level=2)
    if result.workflow_model.prioritized_decisions:
        for item in result.workflow_model.prioritized_decisions:
            execution_rank = item.get("execution_rank", "?")
            decision_id = item.get("decision_id", "unknown-decision")
            priority_score = item.get("priority_score", "?")
            text = item.get("text", item.get("summary", "No decision text available."))
            document.add_paragraph(
                f"Rank {execution_rank} | {decision_id} | score {priority_score} | {text}",
                style="List Bullet",
            )
    document.add_heading("Execution Plan", level=2)
    if result.workflow_model.execution_plan:
        for task in result.workflow_model.execution_plan:
            execution_order = task.get("execution_order", "?")
            title = task.get("title", "Untitled task")
            owner = task.get("owner", "Unassigned")
            priority_score = task.get("priority_score", "?")
            document.add_paragraph(
                f"{execution_order} | {title} | owner {owner} | priority {priority_score}",
                style="List Bullet",
            )
    document.save(path)


def write_pdf_report(result: PipelineResult, path: Path) -> None:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfgen import canvas

    pdf = canvas.Canvas(str(path), pagesize=A4)
    y = 800
    pdf.setFont("Helvetica-Bold", 16)
    pdf.drawString(50, y, "BoardSight Structured Report")
    y -= 30
    pdf.setFont("Helvetica", 11)
    for line in [
        f"Impact score: {result.meeting_scores.impact_score}",
        f"Productivity score: {result.meeting_scores.productivity_score}",
        f"Execution readiness: {result.meeting_scores.execution_readiness}",
        f"Attention: {result.attention_sentiment.overall_attention}",
        f"Sentiment: {result.attention_sentiment.overall_sentiment}",
        f"Conclusion: {result.meeting_scores.meeting_conclusion}",
    ]:
        pdf.drawString(50, y, line[:110])
        y -= 18
    y -= 10
    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawString(50, y, "Decision Moments")
    y -= 20
    pdf.setFont("Helvetica", 10)
    items = result.decision_moments or []
    if not items:
        pdf.drawString(50, y, "No explicit decision moments detected.")
    else:
        for item in items[:10]:
            pdf.drawString(50, y, f"{item.event_id}: {item.text[:95]}")
            y -= 16
            if y < 80:
                pdf.showPage()
                y = 800
    y -= 20
    pdf.setFont("Helvetica-Bold", 13)
    pdf.drawString(50, y, "Execution Plan")
    y -= 18
    pdf.setFont("Helvetica", 10)
    for task in result.workflow_model.execution_plan[:8]:
        execution_order = task.get("execution_order", "?")
        title = str(task.get("title", "Untitled task"))
        pdf.drawString(50, y, f"{execution_order}. {title[:88]}")
        y -= 16
        if y < 80:
            pdf.showPage()
            y = 800
    pdf.save()


def write_excel_report(result: PipelineResult, path: Path) -> None:
    import pandas as pd

    with pd.ExcelWriter(path, engine="xlsxwriter") as writer:
        pd.DataFrame(
            [
                {
                    "impact_score": result.meeting_scores.impact_score,
                    "productivity_score": result.meeting_scores.productivity_score,
                    "execution_readiness": result.meeting_scores.execution_readiness,
                    "overall_attention": result.attention_sentiment.overall_attention,
                    "overall_sentiment": result.attention_sentiment.overall_sentiment,
                }
            ]
        ).to_excel(writer, sheet_name="summary", index=False)
        pd.DataFrame([segment.__dict__ for segment in result.transcript.segments]).to_excel(
            writer, sheet_name="transcript", index=False
        )
        pd.DataFrame([item.__dict__ for item in result.decision_moments]).to_excel(
            writer, sheet_name="decisions", index=False
        )
        pd.DataFrame(result.workflow_model.prioritized_decisions).to_excel(
            writer, sheet_name="prioritized_decisions", index=False
        )
        pd.DataFrame(result.workflow_model.execution_plan).to_excel(
            writer, sheet_name="execution_plan", index=False
        )
        pd.DataFrame([artifact.__dict__ for artifact in result.visual_artifacts]).to_excel(
            writer, sheet_name="visual_artifacts", index=False
        )


def write_summary_image(result: PipelineResult, path: Path) -> None:
    import matplotlib

    matplotlib.use("Agg")
    import matplotlib.pyplot as plt

    fig, ax = plt.subplots(figsize=(6, 4))
    ax.axis("off")
    ax.text(0.05, 0.85, "BoardSight Summary", fontsize=16, weight="bold")
    ax.text(0.05, 0.65, f"Impact: {result.meeting_scores.impact_score}")
    ax.text(0.05, 0.52, f"Productivity: {result.meeting_scores.productivity_score}")
    ax.text(0.05, 0.39, f"Execution readiness: {result.meeting_scores.execution_readiness}")
    ax.text(0.05, 0.26, f"Top decision: {result.workflow_model.workflow_summary.get('top_priority_decision', 'None')}")
    ax.text(0.05, 0.13, f"Tasks queued: {result.workflow_model.workflow_summary.get('total_execution_tasks', 0)}")
    fig.tight_layout()
    fig.savefig(path, dpi=180)
    plt.close(fig)
