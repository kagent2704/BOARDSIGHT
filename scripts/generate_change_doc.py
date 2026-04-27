from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DOCX = ROOT / "docs" / "BoardSight_Changes_and_Key_Code_Snippets.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_paragraph_page_break_before(paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    page_break = OxmlElement("w:pageBreakBefore")
    p_pr.append(page_break)


def add_code_paragraph(document: Document, snippet: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.left_indent = Inches(0.15)
    run = paragraph.add_run(snippet)
    run.font.name = "Consolas"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    run.font.size = Pt(8.5)

    p_pr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "4")
        border.set(qn("w:color"), "D9E2F3")
        borders.append(border)
    p_pr.append(borders)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F7F9FC")
    p_pr.append(shd)


def snippet(path: str, start: int, end: int) -> str:
    file_path = ROOT / path
    lines = file_path.read_text(encoding="utf-8").splitlines()
    selected = []
    for index in range(start, min(end, len(lines)) + 1):
        selected.append(f"{index:>4} {lines[index - 1]}")
    return "\n".join(selected)


def add_heading(document: Document, text: str, level: int = 1, page_break: bool = False) -> None:
    paragraph = document.add_paragraph()
    if page_break:
        set_paragraph_page_break_before(paragraph)
    style = f"Heading {level}"
    paragraph.style = style
    run = paragraph.add_run(text)
    if level == 1:
        run.font.color.rgb = RGBColor(31, 69, 170)


def build_document() -> None:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("BoardSight Change Summary and Key Code Snippets")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(24, 32, 54)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Documented additions across the AI pipeline, Java web app, and CV visibility layer.")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10.5)
    subtitle_run.font.color.rgb = RGBColor(101, 112, 135)

    document.add_paragraph(
        "This document records the major changes added to the BoardSight_CV project during the recent extension work. "
        "Each section includes a short rationale and one or more key code snippets taken from the live codebase."
    )

    add_heading(document, "Change Overview", level=1)
    rows = [
        ("CV Results UI", "Added a dedicated Computer Vision Coverage panel with display-mode, artifact, and timeline views.", "index.html, app.js, styles.css"),
        ("Pipeline Speed", "Improved CPU practicality with lighter default sampling and model caching.", "config.py, speech.py, llm.py"),
        ("Range-based Analysis", "Added fast FFmpeg preprocessing and user-selected start/end analysis windows.", "media.py, service.py, cli.py, PythonPipelineRunner.java"),
        ("Model Integrity", "Kept all model-backed stages active while preserving the no-heuristics policy.", "pipeline.py"),
        ("Slide Content Extraction", "Extended presentation windows to extract best-effort OCR text from slides and documents.", "visual_artifacts.py, models.py, app.js"),
    ]
    for area, outcome, files in rows:
        paragraph = document.add_paragraph(style="List Bullet")
        area_run = paragraph.add_run(f"{area}: ")
        area_run.bold = True
        paragraph.add_run(f"{outcome} Primary files: {files}.")

    sections = [
        {
            "title": "1. Computer Vision Coverage Panel",
            "summary": "A new meeting-detail extension was added to expose the CV pipeline more transparently. The UI now explains frame splitting, speaker-versus-presentation distinctions, artifact breakdowns, and the sampled visual timeline.",
            "snippets": [
                ("Meeting detail panel markup", "java-app/src/main/resources/public/index.html", 167, 181),
                ("CV panel rendering logic", "java-app/src/main/resources/public/app.js", 564, 706),
                ("CV panel styling", "java-app/src/main/resources/public/styles.css", 311, 390),
            ],
        },
        {
            "title": "2. Faster CPU-Friendly Defaults",
            "summary": "The default pipeline profile was tightened to reduce unnecessary visual and workflow sampling while still using the same model-backed stack. This lowers the amount of work done per run, especially on CPU-only machines.",
            "snippets": [
                ("Sampling defaults in AppConfig", "python-ai/boardsight_ai/config.py", 29, 38),
            ],
        },
        {
            "title": "3. Model Caching Improvements",
            "summary": "Repeated model initialization was a major source of delay. Faster-Whisper and the summarization pipeline are now cached so they are loaded once and reused across runs.",
            "snippets": [
                ("Cached Faster-Whisper model", "python-ai/boardsight_ai/providers/speech.py", 17, 27),
                ("Transcription now reuses cached model", "python-ai/boardsight_ai/providers/speech.py", 43, 48),
                ("Cached summarizer pipeline", "python-ai/boardsight_ai/providers/llm.py", 11, 29),
            ],
        },
        {
            "title": "4. User-Selected Time Range and Fast Clip Preprocessing",
            "summary": "Instead of dropping model stages, the speed strategy moved to video preprocessing. Users can now choose a start and end time, and BoardSight quickly clips that portion using FFmpeg stream copy before running the full AI pipeline on the chosen segment.",
            "snippets": [
                ("Fast clip preprocessing helper", "python-ai/boardsight_ai/providers/media.py", 52, 101),
                ("Service-side range handling", "python-ai/boardsight_ai/service.py", 255, 292),
                ("CLI support for start/end seconds", "python-ai/boardsight_ai/cli.py", 24, 50),
                ("Java runner passes range fields through", "java-app/src/main/java/com/boardsight/service/PythonPipelineRunner.java", 29, 123),
                ("Upload form inputs for start/end time", "java-app/src/main/resources/public/index.html", 244, 255),
                ("Frontend parsing and request wiring", "java-app/src/main/resources/public/app.js", 802, 839),
            ],
        },
        {
            "title": "5. Full Deep-Learning Pipeline Preserved",
            "summary": "The range-based flow was deliberately implemented without replacing any model stage with heuristics. The pipeline still runs transcription, speaker modeling, visual artifacts, workflow, decision traces, attention/sentiment, and scoring on the selected input segment.",
            "snippets": [
                ("Pipeline stage flow with analysis_range metadata", "python-ai/boardsight_ai/pipeline.py", 26, 93),
            ],
        },
        {
            "title": "6. Slide and Presentation Content Extraction",
            "summary": "Presentation-like windows now trigger a best-effort OCR extraction pass using TrOCR. The resulting text is saved into each visual artifact and displayed in a new Extracted Slide Content area in the CV panel.",
            "snippets": [
                ("VisualArtifact model now stores content_text", "python-ai/boardsight_ai/models.py", 39, 47),
                ("OCR components and presentation detection", "python-ai/boardsight_ai/features/visual_artifacts.py", 28, 85),
                ("Content text extraction and artifact enrichment", "python-ai/boardsight_ai/features/visual_artifacts.py", 114, 155),
                ("Artifact creation with content_text", "python-ai/boardsight_ai/features/visual_artifacts.py", 194, 213),
                ("UI section for extracted slide content", "java-app/src/main/resources/public/app.js", 697, 705),
                ("Renderer for content artifacts", "java-app/src/main/resources/public/app.js", 1107, 1119),
            ],
        },
    ]

    for index, section_data in enumerate(sections):
        add_heading(document, section_data["title"], level=1, page_break=index > 0)
        document.add_paragraph(section_data["summary"])
        for label, path, start, end in section_data["snippets"]:
            label_paragraph = document.add_paragraph()
            label_run = label_paragraph.add_run(f"{label}  ({path}:{start}-{end})")
            label_run.bold = True
            label_run.font.size = Pt(10.5)
            add_code_paragraph(document, snippet(path, start, end))

    add_heading(document, "Verification Notes", level=1, page_break=True)
    document.add_paragraph(
        "The updated Python files were syntax-checked with py_compile. Java source compilation also succeeded after the changes. "
        "A full jar rebuild may be blocked while the running boardsight.jar process is holding the file lock."
    )

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_document()
